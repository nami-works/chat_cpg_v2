import os
import uuid
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any
import aiofiles
import asyncio
from datetime import datetime

from fastapi import UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update
import pandas as pd
import docx
import PyPDF2
import json
from io import BytesIO

from ..models.knowledge import Document, DocumentChunk, DocumentStatus, DocumentType
from ..models.user import User
from ..core.config import settings
from .usage_service import UsageService


class FileService:
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
        self.usage_service = UsageService()
    
    async def upload_file(
        self,
        file: UploadFile,
        user_id: str,
        knowledge_base_id: str,
        db: AsyncSession,
        tags: Optional[List[str]] = None,
        title: Optional[str] = None
    ) -> Document:
        """
        Upload and process a file.
        """
        # Check file upload limits
        usage_check = await self.usage_service.check_file_upload_limit(user_id, db)
        if not usage_check.allowed:
            raise HTTPException(
                status_code=429,
                detail=f"File upload limit exceeded. {usage_check.message}"
            )
        
        # Validate file
        await self._validate_file(file)
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = self._get_file_extension(file.filename)
        safe_filename = f"{file_id}.{file_extension}"
        
        # Create user directory
        user_dir = self.upload_dir / str(user_id)
        user_dir.mkdir(exist_ok=True)
        
        file_path = user_dir / safe_filename
        
        try:
            # Save file
            content = await file.read()
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(content)
            
            # Get file info
            file_size = len(content)
            file_hash = hashlib.md5(content).hexdigest()
            
            # Create document record
            document = Document(
                id=uuid.UUID(file_id),
                user_id=uuid.UUID(user_id),
                knowledge_base_id=uuid.UUID(knowledge_base_id),
                filename=safe_filename,
                original_filename=file.filename,
                file_type=file_extension.lower(),
                file_size=file_size,
                file_path=str(file_path.relative_to(self.upload_dir)),
                title=title or file.filename,
                status=DocumentStatus.PROCESSING.value,
                processing_started_at=datetime.utcnow(),
                tags=tags or [],
                metadata={
                    "content_hash": file_hash,
                    "mime_type": file.content_type
                }
            )
            
            db.add(document)
            await db.commit()
            await db.refresh(document)
            
            # Increment usage
            await self.usage_service.increment_file_upload_usage(user_id, db)
            
            # Process file asynchronously
            asyncio.create_task(self._process_file_async(document.id, db))
            
            return document
            
        except Exception as e:
            # Clean up file if creation failed
            if file_path.exists():
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")
    
    async def _validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file.
        """
        # Check file size
        if file.size and file.size > settings.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size is {settings.MAX_FILE_SIZE / 1024 / 1024:.1f}MB"
            )
        
        # Check file type
        file_extension = self._get_file_extension(file.filename)
        if file_extension.lower() not in settings.ALLOWED_FILE_TYPES:
            raise HTTPException(
                status_code=415,
                detail=f"File type not supported. Allowed types: {', '.join(settings.ALLOWED_FILE_TYPES)}"
            )
    
    def _get_file_extension(self, filename: str) -> str:
        """
        Get file extension from filename.
        """
        if not filename:
            raise HTTPException(status_code=400, detail="Filename is required")
        
        extension = Path(filename).suffix.lstrip('.')
        if not extension:
            raise HTTPException(status_code=400, detail="File must have an extension")
        
        return extension
    
    async def _process_file_async(self, document_id: uuid.UUID, db: AsyncSession) -> None:
        """
        Process file content asynchronously.
        """
        try:
            # Get document
            result = await db.execute(
                select(Document).where(Document.id == document_id)
            )
            document = result.scalar_one_or_none()
            
            if not document:
                return
            
            # Extract content
            content, metadata = await self._extract_content(document)
            
            # Update document with extracted content
            await db.execute(
                update(Document)
                .where(Document.id == document_id)
                .values(
                    content_preview=content[:500] if content else None,
                    word_count=len(content.split()) if content else 0,
                    metadata={**document.meta, **metadata},
                    processing_completed_at=datetime.utcnow(),
                    status=DocumentStatus.COMPLETED.value
                )
            )
            
            await db.commit()
            
            # Create chunks if content was extracted
            if content:
                await self._create_chunks(document, content, db)
            
        except Exception as e:
            # Update document with error status
            await db.execute(
                update(Document)
                .where(Document.id == document_id)
                .values(
                    status=DocumentStatus.FAILED.value,
                    processing_error=str(e),
                    processing_completed_at=datetime.utcnow()
                )
            )
            await db.commit()
    
    async def _extract_content(self, document: Document) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from file.
        """
        file_path = self.upload_dir / document.file_path
        file_type = document.file_type.lower()
        metadata = {}
        
        if file_type == "pdf":
            return await self._extract_pdf_content(file_path, metadata)
        elif file_type == "docx":
            return await self._extract_docx_content(file_path, metadata)
        elif file_type == "xlsx":
            return await self._extract_xlsx_content(file_path, metadata)
        elif file_type == "csv":
            return await self._extract_csv_content(file_path, metadata)
        elif file_type in ["txt", "md"]:
            return await self._extract_text_content(file_path, metadata)
        elif file_type == "json":
            return await self._extract_json_content(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_content(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = ""
                
                for page_num, page in enumerate(reader.pages):
                    content += f"\n\n--- Page {page_num + 1} ---\n\n"
                    content += page.extract_text()
                
                metadata.update({
                    "total_pages": len(reader.pages),
                    "pdf_info": reader.metadata if reader.metadata else {}
                })
                
                return content.strip(), metadata
        except Exception as e:
            raise ValueError(f"Failed to extract PDF content: {str(e)}")
    
    async def _extract_docx_content(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from DOCX file."""
        try:
            doc = docx.Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            metadata.update({
                "total_paragraphs": len(doc.paragraphs),
                "has_tables": len(doc.tables) > 0,
                "table_count": len(doc.tables)
            })
            
            return content.strip(), metadata
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX content: {str(e)}")
    
    async def _extract_xlsx_content(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from XLSX file."""
        try:
            df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
            content = ""
            
            for sheet_name, sheet_df in df.items():
                content += f"\n\n--- Sheet: {sheet_name} ---\n\n"
                content += sheet_df.to_string(index=False)
            
            metadata.update({
                "sheet_names": list(df.keys()),
                "total_sheets": len(df),
                "total_rows": sum(len(sheet_df) for sheet_df in df.values()),
                "total_columns": sum(len(sheet_df.columns) for sheet_df in df.values())
            })
            
            return content.strip(), metadata
        except Exception as e:
            raise ValueError(f"Failed to extract XLSX content: {str(e)}")
    
    async def _extract_csv_content(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from CSV file."""
        try:
            df = pd.read_csv(file_path)
            content = df.to_string(index=False)
            
            metadata.update({
                "total_rows": len(df),
                "total_columns": len(df.columns),
                "column_names": df.columns.tolist()
            })
            
            return content, metadata
        except Exception as e:
            raise ValueError(f"Failed to extract CSV content: {str(e)}")
    
    async def _extract_text_content(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from text file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            metadata.update({
                "encoding": "utf-8",
                "line_count": len(content.splitlines())
            })
            
            return content, metadata
        except Exception as e:
            raise ValueError(f"Failed to extract text content: {str(e)}")
    
    async def _extract_json_content(self, file_path: Path, metadata: Dict) -> Tuple[str, Dict]:
        """Extract content from JSON file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                json_content = await file.read()
                parsed_json = json.loads(json_content)
            
            # Convert JSON to readable text format
            content = json.dumps(parsed_json, indent=2, ensure_ascii=False)
            
            metadata.update({
                "json_structure": self._analyze_json_structure(parsed_json),
                "encoding": "utf-8"
            })
            
            return content, metadata
        except Exception as e:
            raise ValueError(f"Failed to extract JSON content: {str(e)}")
    
    def _analyze_json_structure(self, json_obj: Any) -> Dict[str, Any]:
        """Analyze JSON structure for metadata."""
        if isinstance(json_obj, dict):
            return {
                "type": "object",
                "keys": list(json_obj.keys()),
                "key_count": len(json_obj)
            }
        elif isinstance(json_obj, list):
            return {
                "type": "array",
                "length": len(json_obj),
                "item_types": list(set(type(item).__name__ for item in json_obj[:10]))  # Sample first 10
            }
        else:
            return {
                "type": type(json_obj).__name__
            }
    
    async def _create_chunks(self, document: Document, content: str, db: AsyncSession) -> None:
        """
        Create text chunks for vector embedding.
        """
        # Simple chunking strategy - can be enhanced later
        chunk_size = 1000  # characters
        chunk_overlap = 200
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(content):
            end = start + chunk_size
            
            # Try to break at word boundary
            if end < len(content):
                while end > start and content[end] not in [' ', '\n', '\t', '.', '!', '?']:
                    end -= 1
            
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                chunk = DocumentChunk(
                    document_id=document.id,
                    user_id=document.user_id,
                    knowledge_base_id=document.knowledge_base_id,
                    chunk_index=chunk_index,
                    content=chunk_content,
                    content_length=len(chunk_content),
                    start_char=start,
                    end_char=end,
                    context_before=content[max(0, start-100):start] if start > 0 else None,
                    context_after=content[end:end+100] if end < len(content) else None
                )
                chunks.append(chunk)
                chunk_index += 1
            
            start = end - chunk_overlap
        
        # Save chunks
        if chunks:
            db.add_all(chunks)
            
            # Update document chunk count
            await db.execute(
                update(Document)
                .where(Document.id == document.id)
                .values(
                    total_chunks=len(chunks),
                    chunks_processed=len(chunks)
                )
            )
            
            await db.commit()
    
    async def delete_file(self, document_id: str, user_id: str, db: AsyncSession) -> bool:
        """
        Delete a file and its associated data.
        """
        # Get document
        result = await db.execute(
            select(Document).where(
                Document.id == uuid.UUID(document_id),
                Document.user_id == uuid.UUID(user_id)
            )
        )
        document = result.scalar_one_or_none()
        
        if not document:
            return False
        
        # Delete physical file
        file_path = self.upload_dir / document.file_path
        if file_path.exists():
            os.remove(file_path)
        
        # Update document status
        await db.execute(
            update(Document)
            .where(Document.id == uuid.UUID(document_id))
            .values(status=DocumentStatus.DELETED.value)
        )
        
        await db.commit()
        return True
    
    async def get_file_content(self, document_id: str, user_id: str, db: AsyncSession) -> Optional[bytes]:
        """
        Get file content for download.
        """
        result = await db.execute(
            select(Document).where(
                Document.id == uuid.UUID(document_id),
                Document.user_id == uuid.UUID(user_id)
            )
        )
        document = result.scalar_one_or_none()
        
        if not document:
            return None
        
        file_path = self.upload_dir / document.file_path
        if not file_path.exists():
            return None
        
        async with aiofiles.open(file_path, 'rb') as f:
            return await f.read() 